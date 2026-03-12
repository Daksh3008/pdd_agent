# document/pdd_generator.py

"""
PDD/BRD Document Generator.
Unified DOCX generator used by both audio and video pipelines.
Supports flowchart insertion (SVG converted to PNG for DOCX, or direct PNG),
screenshot annotation, and detailed steps.
Strips all markdown formatting before inserting into DOCX.
"""

import os
import re
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.config import config


def _strip_markdown_for_docx(text: str) -> str:
    """Remove all markdown formatting before inserting into DOCX."""
    if not text:
        return text
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def _resolve_step_number(step_num, annotated_frames: Dict) -> str:
    """
    Try multiple key formats to find a frame in annotated_frames.
    step_num could be: "2.4.1", 1, "1", etc.
    annotated_frames keys could be: 1, "1", "2.4.1", etc.
    """
    if not annotated_frames:
        return ""

    # Direct lookup
    if step_num in annotated_frames:
        return annotated_frames[step_num]

    # Try string version
    str_num = str(step_num)
    if str_num in annotated_frames:
        return annotated_frames[str_num]

    # Extract trailing integer from "2.4.X"
    if isinstance(step_num, str) and '.' in step_num:
        try:
            num_key = int(step_num.split('.')[-1])
            if num_key in annotated_frames:
                return annotated_frames[num_key]
            if str(num_key) in annotated_frames:
                return annotated_frames[str(num_key)]
        except (ValueError, IndexError):
            pass

    # Try int conversion
    try:
        int_key = int(step_num)
        if int_key in annotated_frames:
            return annotated_frames[int_key]
    except (ValueError, TypeError):
        pass

    return ""


def _convert_svg_to_png(svg_path: str) -> Optional[str]:
    """
    Convert SVG to PNG for DOCX embedding.
    DOCX does not natively support SVG, so we convert.
    Tries multiple methods: cairosvg, Pillow+cairosvg, Inkscape CLI.
    """
    if not svg_path or not os.path.exists(svg_path):
        return None

    png_path = os.path.splitext(svg_path)[0] + '_converted.png'

    # Already a PNG
    if svg_path.lower().endswith('.png'):
        return svg_path

    # Method 1: cairosvg (best quality)
    try:
        import cairosvg
        cairosvg.svg2png(url=svg_path, write_to=png_path, dpi=200)
        if os.path.exists(png_path):
            print(f"    [PDDDoc] SVG→PNG via cairosvg: {png_path}")
            return png_path
    except ImportError:
        pass
    except Exception as e:
        print(f"    [PDDDoc] cairosvg conversion failed: {e}")

    # Method 2: svglib + reportlab
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        drawing = svg2rlg(svg_path)
        if drawing:
            renderPM.drawToFile(drawing, png_path, fmt='PNG', dpi=200)
            if os.path.exists(png_path):
                print(f"    [PDDDoc] SVG→PNG via svglib: {png_path}")
                return png_path
    except ImportError:
        pass
    except Exception as e:
        print(f"    [PDDDoc] svglib conversion failed: {e}")

    # Method 3: Inkscape CLI
    try:
        import subprocess
        result = subprocess.run(
            ['inkscape', svg_path, '--export-type=png',
             f'--export-filename={png_path}', '--export-dpi=200'],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and os.path.exists(png_path):
            print(f"    [PDDDoc] SVG→PNG via Inkscape: {png_path}")
            return png_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception as e:
        print(f"    [PDDDoc] Inkscape conversion failed: {e}")

    # Method 4: Pillow with wand (ImageMagick)
    try:
        from wand.image import Image as WandImage
        with WandImage(filename=svg_path, resolution=200) as img:
            img.format = 'png'
            img.save(filename=png_path)
        if os.path.exists(png_path):
            print(f"    [PDDDoc] SVG→PNG via Wand: {png_path}")
            return png_path
    except ImportError:
        pass
    except Exception as e:
        print(f"    [PDDDoc] Wand conversion failed: {e}")

    print(f"    [PDDDoc] Warning: Could not convert SVG to PNG. "
          f"Install cairosvg (pip install cairosvg) for best results.")
    return None


class PDDGenerator:
    """Unified PDD/BRD document generator."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

    def _setup_margins(self):
        margin = Inches(config.document.margin_inches)
        for section in self.doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

    def _heading(self, text, level=1, color="0D3B66", before=12, after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
        r.font.color.rgb = RGBColor.from_string(color)

    def _para(self, text, after=8):
        text = _strip_markdown_for_docx(text)
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(after)
        return p

    def _table(self, data, header_color="4F81BD", widths=None):
        if not data:
            return
        t = self.doc.add_table(
            rows=len(data), cols=len(data[0]), style="Table Grid"
        )
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        for ri, rd in enumerate(data):
            for ci, ct in enumerate(rd):
                cell = t.cell(ri, ci)
                cell.text = _strip_markdown_for_docx(str(ct))
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                if ri == 0:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), header_color)
                    cell._element.get_or_add_tcPr().append(shd)
        self.doc.add_paragraph()

    def _add_flowchart_image(self, flowchart_path: str, project_name: str):
        """Add flowchart to document with proper sizing. Handles SVG and PNG."""
        if not flowchart_path or not os.path.exists(flowchart_path):
            self._para("[Flowchart to be inserted]")
            return

        # If SVG, convert to PNG for DOCX embedding
        image_path = flowchart_path
        if flowchart_path.lower().endswith('.svg'):
            converted = _convert_svg_to_png(flowchart_path)
            if converted and os.path.exists(converted):
                image_path = converted
            else:
                self._para("[Flowchart generated as SVG — install cairosvg to embed in DOCX]")
                return

        try:
            from PIL import Image
            img = Image.open(image_path)
            img_width, img_height = img.size

            max_width = config.flowchart.max_width_inches
            max_height = 9.0

            aspect_ratio = img_height / img_width
            width = max_width
            height = width * aspect_ratio

            if height > max_height:
                height = max_height
                width = height / aspect_ratio

            self.doc.add_picture(image_path, width=Inches(width))
            self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except ImportError:
            self.doc.add_picture(image_path, width=Inches(6.5))
            self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            print(f"    [PDDDoc] Flowchart image error: {e}")
            self._para("[Flowchart could not be inserted]")
            return

        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(f"Figure 1: {project_name} Process Flow")
        r.italic = True
        r.font.size = Pt(10)

    def _add_multiline_content(self, text: str):
        """Add multi-line text content, handling paragraphs and bullet points."""
        if not text:
            return

        text = _strip_markdown_for_docx(text)

        for para in text.split('\n'):
            para = para.strip()
            if not para:
                continue

            is_bullet = para.startswith(('- ', '• ', '* '))
            if is_bullet:
                para = para.lstrip('-•* ').strip()
                self._para(f"  •  {para}", after=3)
            elif re.match(r'^\d+[\.\)]\s', para):
                self._para(f"  {para}", after=3)
            else:
                self._para(para, after=8)

    def _add_screenshot(self, frame_path: str, step_num: str = ""):
        """Add a screenshot image to the document."""
        if not frame_path or not os.path.exists(frame_path):
            return False

        try:
            self.doc.add_picture(frame_path, width=Inches(5.5))
            last_p = self.doc.paragraphs[-1]
            last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            last_p.paragraph_format.space_after = Pt(10)

            # Add caption
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cap_r = cap_p.add_run(f"Screenshot — Step {step_num}")
            cap_r.italic = True
            cap_r.font.size = Pt(9)
            cap_r.font.color.rgb = RGBColor(120, 120, 120)

            return True
        except Exception as e:
            print(f"    [PDDDoc] Could not add image for step {step_num}: {e}")
            return False

    def generate(
        self,
        project_name: str,
        app_name: str = "",
        document_purpose: str = "",
        overview: str = "",
        justification: str = "",
        as_is: str = "",
        to_be: str = "",
        process_steps: List[Dict] = None,
        input_requirements: List[Dict] = None,
        detailed_steps: List[Dict] = None,
        interface_requirements: List[Dict] = None,
        exception_handling: List[Dict] = None,
        flowchart_path: str = "",
        output_path: str = "PDD.docx",
        annotated_frames: Dict = None
    ) -> str:
        """Generate complete PDD/BRD document."""

        annotated_frames = annotated_frames or {}

        doc_type = config.document.document_type
        doc_type_full = config.document.document_type_full
        steps_header = config.document.process_steps_header
        detailed_header = config.document.detailed_steps_header

        # ===== FRONT PAGE =====
        self.doc.add_paragraph("\n\n\n")
        p = self.doc.add_paragraph()
        r = p.add_run(doc_type_full.upper())
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(13, 59, 102)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph("\n")
        p = self.doc.add_paragraph()
        r = p.add_run(project_name.upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if config.document.confidential:
            self.doc.add_paragraph("\n\n")
            p = self.doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run("CONFIDENTIAL")
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(150, 0, 0)

        self.doc.add_page_break()

        # ===== VERSION INFO =====
        self._heading("Template Version Information", 2, before=0)
        self._table([
            ["Author", "Business Unit", "Classification", "Status", "Version"],
            ["", "", config.document.classification,
             config.document.status, config.document.version]
        ], widths=[1.5, 1.5, 1.3, 1.3, 0.8])

        self.doc.add_page_break()

        # ===== REVIEW & APPROVAL =====
        self._heading("REVIEW & APPROVAL", 1, before=0)
        self._table([
            ["Project Name", project_name, "Project ID", ""],
            ["Document", f"{doc_type_full} ({doc_type})",
             "Version", config.document.version]
        ], widths=[1.5, 2.5, 1.0, 1.0])
        self._table([
            ["Name", "Title / Position", "Date", "Signature"],
            ["", "", "", ""],
            ["", "", "", ""]
        ], widths=[2.0, 2.0, 1.5, 1.5])

        self.doc.add_page_break()

        # ===== CONTENTS =====
        self._heading("Contents", 1, before=0)
        toc = [
            "1  Introduction",
            "    1.1  Purpose of this Document",
            "    1.2  Overview and Objective",
            "    1.3  Business Justification",
            "2  Business Requirements",
            "    2.1  \"As Is\" Process",
            "    2.2  \"To Be\" Process (Automated State)",
            "    2.2.1  Process Flow",
            f"    2.2.2  {steps_header}",
            "    2.3  Input Requirements",
            f"    2.4  {detailed_header}",
            "    2.5  Interface Requirements",
            "3  Exception Handling",
        ]
        for item in toc:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)

        self.doc.add_page_break()

        # ===== 1. INTRODUCTION =====
        self._heading("1  INTRODUCTION", 1, before=0)

        self._heading("1.1  PURPOSE OF THIS DOCUMENT", 2)
        if document_purpose:
            self._add_multiline_content(document_purpose)
        else:
            self._para(
                f"This document defines the requirements for the "
                f"{project_name} automation."
            )

        self._heading("1.2  OVERVIEW AND OBJECTIVE", 2)
        if overview:
            self._add_multiline_content(overview)
        else:
            self._para(
                f"The primary objective is to automate the "
                f"{project_name} process."
            )

        self._heading("1.3  BUSINESS JUSTIFICATION", 2)
        if justification:
            self._add_multiline_content(justification)
        else:
            self._para(f"The {project_name} delivers operational efficiency.")

        self.doc.add_page_break()

        # ===== 2. BUSINESS REQUIREMENTS =====
        self._heading("2  BUSINESS REQUIREMENTS", 1, before=0)

        # 2.1 As-Is
        self._heading("2.1  \"AS IS\" PROCESS", 2)
        if as_is:
            self._add_multiline_content(as_is)
        else:
            self._para("Current manual process to be documented.")

        self.doc.add_page_break()

        # 2.2 To-Be
        self._heading("2.2  \"TO BE\" PROCESS (AUTOMATED STATE)", 2)
        if to_be:
            self._add_multiline_content(to_be)

        # 2.2.1 Process Flow
        self._heading("2.2.1  PROCESS FLOW", 3)
        self._add_flowchart_image(flowchart_path, project_name)
        self.doc.add_paragraph()

        # 2.2.2 Process Steps
        self._heading(f"2.2.2  {steps_header.upper()}", 3)
        if process_steps:
            data = [["SL #", steps_header]]
            for s in process_steps:
                num = s.get("number", "")
                desc = _strip_markdown_for_docx(s.get("description", ""))
                first_sentence = desc.split('.')[0] + '.' if '.' in desc else desc[:100]
                data.append([str(num), first_sentence])
            self._table(data, widths=[0.5, 6.0])
        else:
            self._para("Process steps to be documented.")

        # 2.3 Input Requirements
        self._heading("2.3  INPUT REQUIREMENTS", 2)
        if input_requirements:
            data = [["SL #", "Input Parameter", "Description"]]
            for i, inp in enumerate(input_requirements):
                data.append([
                    str(i + 1),
                    _strip_markdown_for_docx(inp.get("parameter", "")),
                    _strip_markdown_for_docx(inp.get("description", ""))
                ])
            self._table(data, widths=[0.5, 2.0, 4.0])
        else:
            self._table([
                ["SL #", "Input Parameter", "Description"],
                ["1", "", ""],
                ["2", "", ""]
            ], widths=[0.5, 2.0, 4.0])

        self.doc.add_page_break()

        # 2.4 DETAILED PROCESS
        self._heading(f"2.4  {detailed_header.upper()}", 2)

        if detailed_steps:
            screenshots_added = 0

            for step in detailed_steps:
                step_num = step.get("number", 0)
                desc = _strip_markdown_for_docx(step.get("description", ""))

                # Determine display number
                if isinstance(step_num, str) and step_num.startswith("2.4."):
                    display_num = step_num
                else:
                    display_num = f"2.4.{step_num}"

                # Step sub-heading
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(f"{display_num}. {desc}")
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Arial'

                # Operations detected
                ops = step.get("operations_detected", [])
                if ops:
                    op_names = [op.get("display_name", "") for op in ops]
                    op_names = [n for n in op_names if n]
                    if op_names:
                        p = self.doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        r = p.add_run(f"Operations: {', '.join(op_names)}")
                        r.italic = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(80, 80, 150)

                # Find screenshot
                frame_path = _resolve_step_number(step_num, annotated_frames)

                if (not frame_path or not os.path.exists(frame_path)):
                    frame_path = step.get("frame_after_path", "")

                if frame_path and os.path.exists(frame_path):
                    if self._add_screenshot(frame_path, display_num):
                        screenshots_added += 1

            if screenshots_added > 0:
                print(f"    [PDDDoc] Added {screenshots_added} screenshots to Section 2.4")
            else:
                print(f"    [PDDDoc] Warning: No screenshots were added to Section 2.4")
        else:
            self._para("Detailed process steps with screenshots to be documented.")

        self.doc.add_page_break()

        # 2.5 Interface Requirements
        self._heading("2.5  INTERFACE REQUIREMENTS", 2)
        if interface_requirements:
            data = [["SL #", "Interface Requirement", "Description"]]
            for i, app in enumerate(interface_requirements):
                data.append([
                    str(i + 1),
                    _strip_markdown_for_docx(app.get("application", "")),
                    _strip_markdown_for_docx(app.get("purpose", ""))
                ])
            self._table(data, widths=[0.5, 2.5, 3.5])
        else:
            self._table([
                ["SL #", "Interface Requirement", "Description"],
                ["1", "", ""]
            ], widths=[0.5, 2.5, 3.5])

        self.doc.add_page_break()

        # ===== 3. EXCEPTION HANDLING =====
        self._heading("3  EXCEPTION HANDLING", 1, before=0)
        if exception_handling:
            data = [["Exception Scenario", "Handling Action"]]
            for exc in exception_handling:
                data.append([
                    _strip_markdown_for_docx(exc.get("exception", "")),
                    _strip_markdown_for_docx(exc.get("handling", ""))
                ])
            self._table(data, widths=[2.5, 4.5])
        else:
            self._table([
                ["Exception Scenario", "Handling Action"],
                ["Application Login Failure",
                 "Stop execution, log the error, and send a notification."],
                ["Record Not Found",
                 "Log the failure, skip the record, and continue processing."],
            ], widths=[2.5, 4.5])

        # Save
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self.doc.save(output_path)
        print(f"  PDD document saved: {output_path}")
        return output_path

    def append_frames_with_text(
        self,
        doc_path: str,
        frame_text_pairs: List[Tuple[str, str]],
        detailed_steps: List[Dict] = None,
        start_step: int = 1
    ):
        """Append screenshots under Section 2.4 format."""
        if not os.path.exists(doc_path):
            return
        if not frame_text_pairs and not detailed_steps:
            return

        doc = Document(doc_path)

        margin = Inches(config.document.margin_inches)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        detailed_header = config.document.detailed_steps_header

        section_24_idx = None
        section_25_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '2.4' in text and detailed_header.upper()[:10] in text.upper():
                section_24_idx = i
            elif '2.4' in text and 'HIGH LEVEL' in text.upper():
                section_24_idx = i
            elif '2.5' in text and 'INTERFACE' in text.upper():
                section_25_idx = i
                break

        if section_24_idx is not None and section_25_idx is not None:
            for i in range(section_24_idx + 1, section_25_idx):
                if i < len(doc.paragraphs):
                    p = doc.paragraphs[i]
                    for run in p.runs:
                        run.text = ""
                    p.text = ""

        doc.add_page_break()

        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(8)
        r = h.add_run(f"2.4  {detailed_header.upper()}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor.from_string("0D3B66")

        num_frames = len(frame_text_pairs) if frame_text_pairs else 0
        num_detailed = len(detailed_steps) if detailed_steps else 0
        total_steps = max(num_frames, num_detailed)

        for i in range(total_steps):
            step_num = f"2.4.{i + 1}"

            if detailed_steps and i < num_detailed:
                desc = _strip_markdown_for_docx(detailed_steps[i].get("description", ""))
            elif frame_text_pairs and i < num_frames:
                desc = _strip_markdown_for_docx(frame_text_pairs[i][1])
            else:
                desc = "Process step"

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(f"{step_num}. {desc}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = 'Arial'

            if frame_text_pairs and i < num_frames:
                frame_path = frame_text_pairs[i][0]
                if frame_path and os.path.exists(frame_path):
                    try:
                        doc.add_picture(frame_path, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        doc.paragraphs[-1].paragraph_format.space_after = Pt(10)
                    except Exception as e:
                        print(f"    [PDDDoc] Error adding frame {i}: {e}")

        doc.save(doc_path)
        print(
            f"  Appended {total_steps} detailed steps "
            f"({num_frames} with screenshots) to Section 2.4"
        )