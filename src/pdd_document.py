# src/pdd_document.py

"""
PDD/BRD Document Generator.
Technology-neutral, configurable document structure.
Section 2.4 "High Level To Be Detailed Process" with sub-numbered steps and screenshots.
"""

import os
import re
from typing import List, Tuple, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import doc_config


class PDDGenerator:

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

    def _setup_margins(self):
        """Set narrow margins on all sections."""
        margin = Inches(doc_config.margin_inches)
        for section in self.doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

    def _heading(self, text, level=1, color="0D3B66", before=12, after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
        r.font.color.rgb = RGBColor.from_string(color)

    def _para(self, text, after=8):
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(after)
        return p

    def _table(self, data, header_color="4F81BD", widths=None):
        if not data:
            return
        t = self.doc.add_table(rows=len(data), cols=len(data[0]), style="Table Grid")
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        for ri, rd in enumerate(data):
            for ci, ct in enumerate(rd):
                cell = t.cell(ri, ci)
                cell.text = str(ct)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                if ri == 0:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), header_color)
                    cell._element.get_or_add_tcPr().append(shd)
        self.doc.add_paragraph()

    def generate(
        self,
        project_name: str,
        process_summary: str = "",
        inputs_outputs: str = "",
        flowchart_path: str = "",
        document_purpose: str = None,
        applications: List[Dict] = None,
        process_steps: List[str] = None,
        output_path: str = "PDD.docx",
        overview: str = None,
        justification: str = None,
        as_is: str = None,
        to_be: str = None,
        input_requirements: List[Dict] = None,
        exception_handling: List[Dict] = None,
        detailed_steps: List[Dict] = None
    ) -> str:
        """Generate complete PDD/BRD document."""

        doc_type = doc_config.document_type
        doc_type_full = doc_config.document_type_full
        steps_header = doc_config.process_steps_header
        detailed_header = doc_config.detailed_steps_header

        # ===== FRONT PAGE =====
        self.doc.add_paragraph("\n\n\n")
        p = self.doc.add_paragraph()
        r = p.add_run(f"{doc_type_full.upper()}")
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(13, 59, 102)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph("\n")
        p = self.doc.add_paragraph()
        r = p.add_run(project_name.upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if doc_config.confidential:
            self.doc.add_paragraph("\n\n")
            p = self.doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run("CONFIDENTIAL")
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(150, 0, 0)

        self.doc.add_page_break()

        # ===== VERSION INFO =====
        self._heading("Template Version Information", 2, before=0)
        self._table([
            ["Author", "Business Unit", "Classification", "Status", "Version"],
            ["", "", doc_config.classification, doc_config.status, doc_config.version]
        ], widths=[1.5, 1.5, 1.3, 1.3, 0.8])

        self.doc.add_page_break()

        # ===== REVIEW & APPROVAL =====
        self._heading("REVIEW & APPROVAL", 1, before=0)
        self._table([
            ["Project Name", project_name, "Project ID", ""],
            ["Document", f"{doc_type_full} ({doc_type})", "Version", doc_config.version]
        ], widths=[1.5, 2.5, 1.0, 1.0])
        self._table([
            ["Name", "Title / Position", "Date", "Signature"],
            ["", "", "", ""],
            ["", "", "", ""]
        ], widths=[2.0, 2.0, 1.5, 1.5])

        self.doc.add_page_break()

        # ===== CONTENTS =====
        self._heading("Contents", 1, before=0)
        toc = [
            "1  Introduction",
            "    1.1  Purpose of this Document",
            "    1.2  Overview and Objective",
            "    1.3  Business Justification",
            "2  Business Requirements",
            "    2.1  \"As Is\" Process",
            "    2.2  \"To Be\" Process (Automated State)",
            "    2.2.1  Process Flow",
            f"    2.2.2  {steps_header}",
            "    2.3  Input Requirements",
            f"    2.4  {detailed_header}",
            "    2.5  Interface Requirements",
            "3  Exception Handling",
        ]
        for item in toc:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)

        self.doc.add_page_break()

        # ===== 1. INTRODUCTION =====
        self._heading("1  INTRODUCTION", 1, before=0)

        self._heading("1.1  PURPOSE OF THIS DOCUMENT", 2)
        if document_purpose:
            for para in document_purpose.split('\n\n'):
                if para.strip():
                    self._para(para.strip())
        else:
            self._para(
                f"This document defines the requirements for the "
                f"{project_name} automation."
            )

        self._heading("1.2  OVERVIEW AND OBJECTIVE", 2)
        if overview:
            for line in overview.split('\n'):
                line = line.strip()
                if line:
                    self._para(
                        line,
                        after=3 if line.startswith(('•', '-', '*')) else 8
                    )
        else:
            self._para(
                f"The primary objective is to automate the "
                f"{project_name} process."
            )

        self._heading("1.3  BUSINESS JUSTIFICATION", 2)
        if justification:
            for line in justification.split('\n'):
                if line.strip():
                    self._para(line.strip())
        else:
            self._para(
                f"The {project_name} delivers operational efficiency."
            )

        self.doc.add_page_break()

        # ===== 2. BUSINESS REQUIREMENTS =====
        self._heading("2  BUSINESS REQUIREMENTS", 1, before=0)

        # 2.1 As-Is
        self._heading("2.1  \"AS IS\" PROCESS", 2)
        if as_is:
            for line in as_is.split('\n'):
                if line.strip():
                    self._para(line.strip())
        else:
            self._para("Current manual process to be documented.")

        self.doc.add_page_break()

        # 2.2 To-Be
        self._heading("2.2  \"TO BE\" PROCESS (AUTOMATED STATE)", 2)
        tb = to_be or process_summary
        if tb:
            for para in tb.split('\n\n'):
                if para.strip():
                    self._para(para.strip())

        # 2.2.1 Process Flow
        self._heading("2.2.1  PROCESS FLOW", 3)
        if flowchart_path and os.path.exists(flowchart_path):
            self.doc.add_picture(flowchart_path, width=Inches(6.5))
            self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run(f"Figure 1: {project_name} Process Flow")
            r.italic = True
            r.font.size = Pt(10)
        else:
            self._para("[Flowchart to be inserted]")

        self.doc.add_paragraph()

        # 2.2.2 Process Steps
        self._heading(f"2.2.2  {steps_header.upper()}", 3)
        if process_steps:
            data = [["SL #", steps_header]]
            for i, step in enumerate(process_steps):
                data.append([
                    str(i + 1),
                    step.replace('[DECISION]', '').strip()
                ])
            self._table(data, widths=[0.5, 6.0])
        else:
            self._para("Process steps to be documented.")

        # 2.3 Input Requirements
        self._heading("2.3  INPUT REQUIREMENTS", 2)
        if input_requirements:
            data = [["SL #", "Input Parameter", "Description"]]
            for i, inp in enumerate(input_requirements):
                data.append([
                    str(i + 1),
                    inp.get("parameter", ""),
                    inp.get("description", "")
                ])
            self._table(data, widths=[0.5, 2.0, 4.0])
        else:
            self._table([
                ["SL #", "Input Parameter", "Description"],
                ["1", "", ""],
                ["2", "", ""]
            ], widths=[0.5, 2.0, 4.0])

        self.doc.add_page_break()

        # 2.4 HIGH LEVEL TO BE DETAILED PROCESS
        self._heading(f"2.4  {detailed_header.upper()}", 2)

        if detailed_steps:
            for ds in detailed_steps:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(f"{ds['number']}. {ds['description']}")
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Arial'
        else:
            self._para(
                "Detailed process steps with screenshots to be documented."
            )

        self.doc.add_page_break()

        # 2.5 Interface Requirements
        self._heading("2.5  INTERFACE REQUIREMENTS", 2)
        if applications:
            data = [["SL #", "Interface Requirement", "Description"]]
            for i, app in enumerate(applications):
                data.append([
                    str(i + 1),
                    app.get("application", ""),
                    app.get("purpose", "")
                ])
            self._table(data, widths=[0.5, 2.5, 3.5])
        else:
            self._table([
                ["SL #", "Interface Requirement", "Description"],
                ["1", "", ""]
            ], widths=[0.5, 2.5, 3.5])

        self.doc.add_page_break()

        # ===== 3. EXCEPTION HANDLING =====
        self._heading("3  EXCEPTION HANDLING", 1, before=0)
        if exception_handling:
            data = [["Exception Scenario", "Handling Action"]]
            for exc in exception_handling:
                data.append([
                    exc.get("exception", ""),
                    exc.get("handling", "")
                ])
            self._table(data, widths=[2.5, 4.5])
        else:
            self._table([
                ["Exception Scenario", "Handling Action"],
                ["Application Login Failure",
                 "Stop execution, log error, notify support."],
                ["Record Not Found",
                 "Log failure, skip, continue processing."],
                ["Processing Error",
                 "Log error, mark failed, continue with others."],
                ["System Exception",
                 "Capture in error log for audit."]
            ], widths=[2.5, 4.5])

        self.doc.save(output_path)
        print(f"Document saved: {output_path}")
        return output_path

    def append_frames_with_text(
        self,
        doc_path: str,
        frame_text_pairs: List[Tuple[str, str]],
        detailed_steps: List[Dict] = None,
        start_step: int = 1
    ):
        """
        Append screenshots under Section 2.4 format.
        Each frame gets a 2.4.x sub-number with bold description
        and screenshot below.
        """
        if not os.path.exists(doc_path):
            return
        if not frame_text_pairs and not detailed_steps:
            return

        doc = Document(doc_path)

        # Apply narrow margins to loaded document
        margin = Inches(doc_config.margin_inches)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        detailed_header = doc_config.detailed_steps_header

        # Find where Section 2.4 heading is
        section_24_idx = None
        section_25_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '2.4' in text and detailed_header.upper()[:10] in text.upper():
                section_24_idx = i
            elif '2.4' in text and 'HIGH LEVEL' in text.upper():
                section_24_idx = i
            elif '2.5' in text and 'INTERFACE' in text.upper():
                section_25_idx = i
                break

        if section_24_idx is not None and section_25_idx is not None:
            # Clear existing placeholder content between 2.4 and 2.5
            for i in range(section_24_idx + 1, section_25_idx):
                if i < len(doc.paragraphs):
                    p = doc.paragraphs[i]
                    for run in p.runs:
                        run.text = ""
                    p.text = ""

        # Append detailed steps with screenshots at end of document
        doc.add_page_break()

        # Re-add the Section 2.4 heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(8)
        r = h.add_run(f"2.4  {detailed_header.upper()}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor.from_string("0D3B66")

        # Determine total steps to render
        num_frames = len(frame_text_pairs) if frame_text_pairs else 0
        num_detailed = len(detailed_steps) if detailed_steps else 0
        total_steps = max(num_frames, num_detailed)

        for i in range(total_steps):
            step_num = f"2.4.{i + 1}"

            # Get description
            if detailed_steps and i < num_detailed:
                desc = detailed_steps[i].get("description", "")
            elif frame_text_pairs and i < num_frames:
                desc = frame_text_pairs[i][1]
            else:
                desc = "Process step"

            # Step sub-heading
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(f"{step_num}. {desc}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = 'Arial'

            # Screenshot (if available for this step)
            if frame_text_pairs and i < num_frames:
                frame_path = frame_text_pairs[i][0]
                if os.path.exists(frame_path):
                    doc.add_picture(frame_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

        doc.save(doc_path)
        print(
            f"Appended {total_steps} detailed steps "
            f"({num_frames} with screenshots) to Section 2.4"
        )