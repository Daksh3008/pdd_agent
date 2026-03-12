# pdd_no_audio/pdd_document.py

"""
PDD Document Generator for Silent Screen Recording.
High-quality flowchart handling — fits to page without blurring.
"""

import os
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pdd_no_audio.config import doc_config, flowchart_config


class PDDGenerator:

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

    def _setup_margins(self):
        margin = Inches(doc_config.margin_inches)
        for section in self.doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

    def _heading(self, text, level=1, color="0D3B66", before=12, after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
        r.font.color.rgb = RGBColor.from_string(color)

    def _para(self, text, after=8):
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(after)
        return p

    def _table(self, data, header_color="4F81BD", widths=None):
        if not data:
            return
        t = self.doc.add_table(rows=len(data), cols=len(data[0]), style="Table Grid")
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        for ri, rd in enumerate(data):
            for ci, ct in enumerate(rd):
                cell = t.cell(ri, ci)
                cell.text = str(ct)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                if ri == 0:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), header_color)
                    cell._element.get_or_add_tcPr().append(shd)
        self.doc.add_paragraph()

    def _add_flowchart_image(self, flowchart_path: str, project_name: str):
        """
        Add flowchart to document with proper sizing.
        Calculates aspect ratio to fit within page without distortion.
        """
        if not flowchart_path or not os.path.exists(flowchart_path):
            self._para("[Flowchart to be inserted]")
            return

        try:
            from PIL import Image
            img = Image.open(flowchart_path)
            img_width, img_height = img.size

            # Available page width (page minus margins)
            max_width = flowchart_config.max_width_inches
            max_height = 9.0  # Leave room for caption

            # Calculate scaling
            aspect_ratio = img_height / img_width
            width = max_width
            height = width * aspect_ratio

            # If too tall, scale by height instead
            if height > max_height:
                height = max_height
                width = height / aspect_ratio

            self.doc.add_picture(flowchart_path, width=Inches(width))
            self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except ImportError:
            # PIL not available — use fixed width
            self.doc.add_picture(flowchart_path, width=Inches(6.5))
            self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            print(f"    [PDDDoc] Flowchart image error: {e}")
            self._para("[Flowchart could not be inserted]")
            return

        # Caption
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(f"Figure 1: {project_name} Process Flow")
        r.italic = True
        r.font.size = Pt(10)

    def generate(
        self,
        project_name: str,
        app_name: str = "",
        document_purpose: str = "",
        overview: str = "",
        justification: str = "",
        as_is: str = "",
        to_be: str = "",
        process_steps: List[Dict] = None,
        input_requirements: List[Dict] = None,
        detailed_steps: List[Dict] = None,
        interface_requirements: List[Dict] = None,
        exception_handling: List[Dict] = None,
        flowchart_path: str = "",
        output_path: str = "PDD.docx",
        annotated_frames: Dict[int, str] = None
    ) -> str:

        doc_type = doc_config.document_type
        doc_type_full = doc_config.document_type_full
        steps_header = doc_config.process_steps_header
        detailed_header = doc_config.detailed_steps_header

        # ===== FRONT PAGE =====
        self.doc.add_paragraph("\n\n\n")
        p = self.doc.add_paragraph()
        r = p.add_run(doc_type_full.upper())
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(13, 59, 102)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph("\n")
        p = self.doc.add_paragraph()
        r = p.add_run(project_name.upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if doc_config.confidential:
            self.doc.add_paragraph("\n\n")
            p = self.doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run("CONFIDENTIAL")
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(150, 0, 0)

        self.doc.add_page_break()

        # ===== VERSION INFO =====
        self._heading("Template Version Information", 2, before=0)
        self._table([
            ["Author", "Business Unit", "Classification", "Status", "Version"],
            ["", "", doc_config.classification, doc_config.status, doc_config.version]
        ], widths=[1.5, 1.5, 1.3, 1.3, 0.8])

        self.doc.add_page_break()

        # ===== REVIEW & APPROVAL =====
        self._heading("REVIEW & APPROVAL", 1, before=0)
        self._table([
            ["Project Name", project_name, "Project ID", ""],
            ["Document", f"{doc_type_full} ({doc_type})", "Version", doc_config.version]
        ], widths=[1.5, 2.5, 1.0, 1.0])
        self._table([
            ["Name", "Title / Position", "Date", "Signature"],
            ["", "", "", ""], ["", "", "", ""]
        ], widths=[2.0, 2.0, 1.5, 1.5])

        self.doc.add_page_break()

        # ===== CONTENTS =====
        self._heading("Contents", 1, before=0)
        toc = [
            "1  Introduction",
            "    1.1  Purpose of this Document",
            "    1.2  Overview and Objective",
            "    1.3  Business Justification",
            "2  Business Requirements",
            "    2.1  \"As Is\" Process",
            "    2.2  \"To Be\" Process (Automated State)",
            "    2.2.1  Process Flow",
            f"    2.2.2  {steps_header}",
            "    2.3  Input Requirements",
            f"    2.4  {detailed_header}",
            "    2.5  Interface Requirements",
            "3  Exception Handling",
        ]
        for item in toc:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)

        self.doc.add_page_break()

        # ===== 1. INTRODUCTION =====
        self._heading("1  INTRODUCTION", 1, before=0)

        self._heading("1.1  PURPOSE OF THIS DOCUMENT", 2)
        if document_purpose:
            for para in document_purpose.split('\n\n'):
                if para.strip():
                    self._para(para.strip())
        else:
            self._para(f"This document defines the requirements for the {project_name} automation.")

        self._heading("1.2  OVERVIEW AND OBJECTIVE", 2)
        if overview:
            for line in overview.split('\n'):
                line = line.strip()
                if line:
                    self._para(line, after=3 if line.startswith(('•', '-', '*')) else 8)
        else:
            self._para(f"The primary objective is to automate the {project_name} process.")

        self._heading("1.3  BUSINESS JUSTIFICATION", 2)
        if justification:
            for line in justification.split('\n'):
                if line.strip():
                    self._para(line.strip())
        else:
            self._para(f"The {project_name} delivers operational efficiency.")

        self.doc.add_page_break()

        # ===== 2. BUSINESS REQUIREMENTS =====
        self._heading("2  BUSINESS REQUIREMENTS", 1, before=0)

        self._heading("2.1  \"AS IS\" PROCESS", 2)
        if as_is:
            for line in as_is.split('\n'):
                if line.strip():
                    self._para(line.strip())
        else:
            self._para("Current manual process to be documented.")

        self.doc.add_page_break()

        self._heading("2.2  \"TO BE\" PROCESS (AUTOMATED STATE)", 2)
        if to_be:
            for para in to_be.split('\n\n'):
                if para.strip():
                    self._para(para.strip())

        # 2.2.1 Process Flow — using new method
        self._heading("2.2.1  PROCESS FLOW", 3)
        self._add_flowchart_image(flowchart_path, project_name)

        self.doc.add_paragraph()

        # 2.2.2 Process Steps
        self._heading(f"2.2.2  {steps_header.upper()}", 3)
        if process_steps:
            data = [["SL #", steps_header]]
            for s in process_steps:
                desc = s["description"]
                first_sentence = desc.split('.')[0] + '.' if '.' in desc else desc[:100]
                data.append([str(s["number"]), first_sentence])
            self._table(data, widths=[0.5, 6.0])
        else:
            self._para("Process steps to be documented.")

        # 2.3 Input Requirements
        self._heading("2.3  INPUT REQUIREMENTS", 2)
        if input_requirements:
            data = [["SL #", "Input Parameter", "Description"]]
            for i, inp in enumerate(input_requirements):
                data.append([str(i + 1), inp.get("parameter", ""), inp.get("description", "")])
            self._table(data, widths=[0.5, 2.0, 4.0])
        else:
            self._table([
                ["SL #", "Input Parameter", "Description"],
                ["1", "", ""], ["2", "", ""]
            ], widths=[0.5, 2.0, 4.0])

        self.doc.add_page_break()

        # 2.4 DETAILED PROCESS
        self._heading(f"2.4  {detailed_header.upper()}", 2)

        if detailed_steps:
            annotated_frames = annotated_frames or {}
            for step in detailed_steps:
                step_num = step["number"]
                desc = step["description"]

                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(f"2.4.{step_num}. {desc}")
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Arial'

                ops = step.get("operations_detected", [])
                if ops:
                    op_names = [op["display_name"] for op in ops]
                    p = self.doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(f"Operations: {', '.join(op_names)}")
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(80, 80, 150)

                frame_path = annotated_frames.get(step_num, "")
                if not frame_path or not os.path.exists(frame_path):
                    frame_path = step.get("frame_after_path", "")

                if frame_path and os.path.exists(frame_path):
                    try:
                        self.doc.add_picture(frame_path, width=Inches(5.5))
                        self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        self.doc.paragraphs[-1].paragraph_format.space_after = Pt(10)
                    except Exception as e:
                        print(f"    [PDDDoc] Could not add image for step {step_num}: {e}")
        else:
            self._para("Detailed process steps with screenshots to be documented.")

        self.doc.add_page_break()

        # 2.5 Interface Requirements
        self._heading("2.5  INTERFACE REQUIREMENTS", 2)
        if interface_requirements:
            data = [["SL #", "Interface Requirement", "Description"]]
            for i, app in enumerate(interface_requirements):
                data.append([str(i + 1), app.get("application", ""), app.get("purpose", "")])
            self._table(data, widths=[0.5, 2.5, 3.5])
        else:
            self._table([
                ["SL #", "Interface Requirement", "Description"],
                ["1", "", ""]
            ], widths=[0.5, 2.5, 3.5])

        self.doc.add_page_break()

        # ===== 3. EXCEPTION HANDLING =====
        self._heading("3  EXCEPTION HANDLING", 1, before=0)
        if exception_handling:
            data = [["Exception Scenario", "Handling Action"]]
            for exc in exception_handling:
                data.append([exc.get("exception", ""), exc.get("handling", "")])
            self._table(data, widths=[2.5, 4.5])
        else:
            self._table([
                ["Exception Scenario", "Handling Action"],
                ["Application Login Failure", "Stop execution, log error, notify support."],
                ["Record Not Found", "Log failure, skip, continue processing."],
            ], widths=[2.5, 4.5])

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self.doc.save(output_path)
        print(f"  📄 PDD document saved: {output_path}")
        return output_path